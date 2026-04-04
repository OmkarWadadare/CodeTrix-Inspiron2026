"""
rebuilder.py

rebuild_pdf_translated(json_path, translations, output_pdf)
rebuild_docx_translated(json_path, translations, output_docx)

`translations` is a dict mapping segment_id → translated text string.
Any segment_id not present in `translations` falls back to its original text.
"""

import json
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image

from docx import Document
from docx.shared import Pt


# ─────────────────────────────────────────────
#  PDF REBUILDER
# ─────────────────────────────────────────────

def rebuild_pdf_translated(json_path: str, translations: dict, output_pdf: str) -> None:
    """
    Rebuild PDF from segment-based JSON using helper rendering pipeline.

    Parameters
    ----------
    json_path : str
        Path to pdf.json (segment format)
    translations : dict
        {segment_id: translated_text}
    output_pdf : str
        Output PDF path
    """
    json_path = Path(json_path)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    c = canvas.Canvas(str(output_pdf))

    pages_meta   = data["pages"]
    elements     = data["elements"]
    text_content = data["text_content"]
    assets       = data.get("assets", {})

    num_pages = len(pages_meta)

    for i in range(num_pages):
        page_meta = pages_meta[str(i)] if isinstance(pages_meta, dict) else pages_meta[i]
        pw = page_meta["width"]
        ph = page_meta["height"]

        c.setPageSize((pw, ph))

        # ─────────────────────────────────────────────
        # Build page_data (convert to helper format)
        # ─────────────────────────────────────────────
        page_data = {
            "text_blocks": [],
            "tables": [],
            "images": [],
            "links": []
        }

        for el in elements:
            if el["page"] != i:
                continue

            # ── TEXT → text_blocks ───────────────────
            if el["type"] == "text":
                raw_text = translations.get(
                    el["id"],
                    text_content.get(el["id"], "")
                ).strip()

                if not raw_text:
                    continue

                x0, y0, x1, y1 = el["bbox"]
                flags = el.get("flags", 0)

                block = {
                    "text": raw_text,
                    "x0": x0,
                    "y0": y0,
                    "x1": x1,
                    "y1": y1,
                    "size": el.get("size", 10),
                    "bold": bool(flags & 2),
                    "italic": bool(flags & 1),
                    "color": el.get("color"),
                    "role": "body"
                }

                page_data["text_blocks"].append(block)

            # ── IMAGE → images ───────────────────────
            elif el["type"] == "image":
                asset = assets.get(el["id"])
                if not asset:
                    continue

                page_data["images"].append({
                    "file": asset["path"],
                    "bbox": {
                        "x0": el["bbox"][0],
                        "y0": el["bbox"][1],
                        "x1": el["bbox"][2],
                        "y1": el["bbox"][3],
                    }
                })

            # ── LINK (if exists in your format) ──────
            elif el["type"] == "link":
                page_data["links"].append({
                    "bbox": {
                        "x0": el["bbox"][0],
                        "y0": el["bbox"][1],
                        "x1": el["bbox"][2],
                        "y1": el["bbox"][3],
                    },
                    "url": el.get("url")
                })

            # ── TABLE (if your JSON has it) ──────────
            elif el["type"] == "table":
                page_data["tables"].append(el)

        # ─────────────────────────────────────────────
        # Rendering using helpers
        # ─────────────────────────────────────────────

        table_bboxes = [
            tbl["bbox"]
            for tbl in page_data.get("tables", [])
            if tbl.get("bbox")
        ]

        # TEXT (skip ones inside tables)
        for block in page_data.get("text_blocks", []):
            if _block_inside_table(block, table_bboxes):
                continue
            _draw_text_block(c, block, ph)

        # TABLES
        for table in page_data.get("tables", []):
            _draw_table(c, table, ph)

        # IMAGES
        for img_info in page_data.get("images", []):
            _draw_image(c, img_info, json_path.parent, ph)

        # LINKS
        for link in page_data.get("links", []):
            _draw_link(c, link)

        c.showPage()

    c.save()
    print(f"[rebuild_pdf_translated] → {output_pdf}")

def _block_inside_table(block: dict, table_bboxes: list) -> bool:
    """
    Return True if the block's centre point falls within any table bounding box.
    Using the centre (rather than full overlap) handles blocks that slightly
    straddle a table border due to coordinate rounding.
    """
    cx = (block["x0"] + block["x1"]) / 2
    cy = (block["y0"] + block["y1"]) / 2
    for bbox in table_bboxes:
        if bbox["x0"] <= cx <= bbox["x1"] and bbox["y0"] <= cy <= bbox["y1"]:
            return True
    return False


def _draw_text_block(c: canvas.Canvas, block: dict, page_height: float) -> None:
    """Render a text block onto the ReportLab canvas."""
    text  = block.get("text", "").strip()
    if not text:
        return

    x0   = block["x0"]
    y0   = block["y0"]   # already in PDF coords (bottom-left origin)
    size = block.get("size", 11)
    bold   = block.get("bold", False)
    italic = block.get("italic", False)
    color_hex = block.get("color")

    # Font selection
    font_name = _select_reportlab_font(bold, italic)
    try:
        c.setFont(font_name, size)
    except Exception:
        c.setFont("Helvetica", size)

    # Color
    if color_hex:
        try:
            c.setFillColor(HexColor(color_hex))
        except Exception:
            c.setFillColor(colors.black)
    else:
        c.setFillColor(colors.black)

    # Role-based decorations
    role = block.get("role", "body")
    if role in ("h1", "h2", "h3"):
        # Underline headings
        tw = c.stringWidth(text, font_name, size)
        c.line(x0, y0 - 1, x0 + tw, y0 - 1)

    # Render text
    c.drawString(x0, y0, text)

    # Reset color
    c.setFillColor(colors.black)


def _select_reportlab_font(bold: bool, italic: bool) -> str:
    """Map bold/italic flags to a built-in ReportLab font name."""
    if bold and italic:
        return "Helvetica-BoldOblique"
    if bold:
        return "Helvetica-Bold"
    if italic:
        return "Helvetica-Oblique"
    return "Helvetica"


def _draw_table(c: canvas.Canvas, table: dict, page_height: float) -> None:
    rows = table.get("rows", [])
    bbox = table.get("bbox", {})
    if not rows or not bbox:
        return

    # ── ORIGINAL TABLE SIZE (FIXED) ─────────────────────────────────────
    x0, y0, x1, y1 = bbox["x0"], bbox["y0"], bbox["x1"], bbox["y1"]
    table_w = x1 - x0
    table_h = y1 - y0

    num_rows = len(rows)
    num_cols = max((len(r) for r in rows), default=1)

    col_w = table_w / num_cols
    row_h = table_h / num_rows

    CELL_PAD = 4

    # Normalize rows
    norm_rows = [list(r) + [""] * (num_cols - len(r)) for r in rows]

    # ── Step 1: Find MAX font size that fits ALL cells ───────────────────
    def fits(font_size):
        for i, row in enumerate(norm_rows):
            font = "Helvetica-Bold" if i == 0 else "Helvetica"
            c.setFont(font, font_size)

            for j, cell in enumerate(row):
                text = str(cell)
                lines = text.split("\n")

                # vertical check
                needed_h = len(lines) * (font_size + 2)
                if needed_h > (row_h - 2 * CELL_PAD):
                    return False

                # horizontal check
                for line in lines:
                    if c.stringWidth(line, font, font_size) > (col_w - 2 * CELL_PAD):
                        return False
        return True

    # Binary search font size
    low, high = 4, 14
    best_size = 6

    while low <= high:
        mid = (low + high) // 2
        if fits(mid):
            best_size = mid
            low = mid + 1
        else:
            high = mid - 1

    FONT_SIZE = best_size
    LINE_H = FONT_SIZE + 2

    # ── Header shading ──────────────────────────────────────────────────
    c.setFillColorRGB(0.85, 0.85, 0.85)
    c.rect(x0, y1 - row_h, table_w, row_h, fill=1, stroke=0)
    c.setFillColor(colors.black)

    # ── Grid ────────────────────────────────────────────────────────────
    c.setStrokeColor(colors.black)
    c.setLineWidth(0.5)

    # Horizontal lines
    for i in range(num_rows + 1):
        y = y1 - i * row_h
        c.line(x0, y, x1, y)

    # Vertical lines
    for j in range(num_cols + 1):
        x = x0 + j * col_w
        c.line(x, y0, x, y1)

    # ── Draw text (FIT INSIDE CELLS) ────────────────────────────────────
    for i, row in enumerate(norm_rows):
        font = "Helvetica-Bold" if i == 0 else "Helvetica"
        c.setFont(font, FONT_SIZE)

        for j, cell in enumerate(row):
            text = str(cell)
            lines = text.split("\n")

            tx = x0 + j * col_w + CELL_PAD
            ty = y1 - i * row_h - CELL_PAD - LINE_H

            for line in lines:
                c.drawString(tx, ty, line)
                ty -= LINE_H

    # Reset
    c.setFont("Helvetica", 11)
    c.setFillColor(colors.black)

def _draw_image(c: canvas.Canvas, img_info: dict,
                data_dir: Path, page_height: float) -> None:
    """Place an extracted image onto the canvas at the stored position."""
    rel_path = img_info.get("file")
    if not rel_path:
        return

    img_path = data_dir / rel_path
    if not img_path.exists():
        return

    bbox = img_info.get("bbox", {})
    x0   = bbox.get("x0", 0)
    y0   = bbox.get("y0", 0)   # bottom-left in PDF coords
    x1   = bbox.get("x1", x0 + 100)
    y1   = bbox.get("y1", y0 + 100)
    w    = max(x1 - x0, 1)
    h    = max(y1 - y0, 1)

    try:
        pil_img = Image.open(img_path)
        img_reader = ImageReader(pil_img)
        c.drawImage(img_reader, x0, y0, width=w, height=h,
                    preserveAspectRatio=True, mask="auto")
    except Exception as e:
        # Draw a placeholder box if image can't be rendered
        c.setStrokeColor(colors.grey)
        c.setFillColorRGB(0.95, 0.95, 0.95)
        c.rect(x0, y0, w, h, fill=1)
        c.setFillColor(colors.grey)
        c.setFont("Helvetica", 8)
        c.drawCentredString(x0 + w / 2, y0 + h / 2, "[image]")
        c.setFillColor(colors.black)

def _draw_link(c: canvas.Canvas, link: dict) -> None:
    bbox = link.get("bbox", {})
    url  = link.get("url")

    if not bbox or not url:
        return

    x0 = bbox["x0"]
    y0 = bbox["y0"]
    x1 = bbox["x1"]
    y1 = bbox["y1"]

    # Invisible clickable area
    c.linkURL(
        url,
        (x0, y0, x1, y1),
        relative=0,
        thickness=0
    )

# ─────────────────────────────────────────────
#  DOCX REBUILDER
# ─────────────────────────────────────────────

def rebuild_docx_translated(json_path: str, translations: dict, output_docx: str) -> None:
    """
    Rebuild a translated DOCX from a segments.json produced by extract_docx_segments().

    Parameters
    ----------
    json_path    : str   Path to segments.json
    translations : dict  {segment_id: translated_text, ...}
    output_docx  : str   Output DOCX path
    """
    json_path = Path(json_path)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    seg_map = {s["segment_id"]: s for s in data["segments"]}
    assets  = data.get("assets", {})

    doc = Document()

    for elem in data["elements"]:
        etype = elem["type"]

        if etype == "paragraph":
            p = doc.add_paragraph()
            try:
                p.style = elem["style"]
            except KeyError:
                p.style = "Normal"

            for seg_id in elem.get("segment_ids", []):
                if seg_id is None:
                    # Preserve empty run as blank
                    p.add_run("")
                    continue
                seg = seg_map.get(seg_id)
                if seg is None:
                    continue
                text = translations.get(seg_id, seg["text"])
                run = p.add_run(text)
                run.bold      = seg.get("bold", False)
                run.italic    = seg.get("italic", False)
                run.underline = seg.get("underline", False)
                fn = seg.get("font_name") or "Calibri"
                run.font.name = fn
                fs = seg.get("font_size")
                if fs:
                    run.font.size = Pt(fs)

            # Inline images attached to this paragraph
            for img_ref in elem.get("image_refs", []):
                img_path = assets.get(img_ref)
                if img_path and Path(img_path).exists():
                    try:
                        doc.add_picture(img_path)
                    except Exception:
                        pass

        elif etype == "table":
            rows = elem["rows"]
            cols = elem["cols"]
            if rows == 0 or cols == 0:
                continue

            table = doc.add_table(rows=rows, cols=cols)

            for r_idx, row_data in enumerate(elem.get("cells", [])):
                for c_idx, cell_data in enumerate(row_data):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]

                    style = cell_data.get("style")
                    if style:
                        try:
                            p.style = style
                        except KeyError:
                            pass

                    for seg_id in cell_data.get("segment_ids", []):
                        if seg_id is None:
                            p.add_run("")
                            continue
                        seg = seg_map.get(seg_id)
                        if seg is None:
                            continue
                        text = translations.get(seg_id, seg["text"])
                        run = p.add_run(text)
                        run.bold      = seg.get("bold", False)
                        run.italic    = seg.get("italic", False)
                        run.underline = seg.get("underline", False)
                        fn = seg.get("font_name") or "Calibri"
                        run.font.name = fn
                        fs = seg.get("font_size")
                        if fs:
                            run.font.size = Pt(fs)

    doc.save(output_docx)
    print(f"[rebuild_docx_translated] → {output_docx}")