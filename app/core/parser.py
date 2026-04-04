# app/core/parser.py
import fitz
import os
import json
from docx import Document

def extract_pdf_segments(input_pdf: str, output_dir: str) -> dict:
    """Extract PDF → structured segments (INCLUDING TABLES)."""
    import fitz
    import os
    import json

    os.makedirs(output_dir, exist_ok=True)
    img_dir = os.path.join(output_dir, "images")
    os.makedirs(img_dir, exist_ok=True)

    doc = fitz.open(input_pdf)

    elements, text_content, assets, pages_meta = [], {}, {}, {}
    text_id = img_id = table_id = 0

    for page_num, page in enumerate(doc):
        pages_meta[page_num] = {
            "width": page.rect.width,
            "height": page.rect.height
        }

        # ─────────────────────────────────────────────
        # 1. DETECT TABLES FIRST
        # ─────────────────────────────────────────────
        tables = page.find_tables()

        table_bboxes = []

        if tables:
            for tbl in tables:
                bbox = tbl.bbox
                table_bboxes.append(bbox)

                tid = f"tbl_{table_id}"

                # Extract cell text matrix
                table_data = tbl.extract()  # list of rows

                elements.append({
                    "type": "table",
                    "id": tid,
                    "page": page_num,
                    "bbox": bbox,
                    "rows": table_data  # <-- IMPORTANT
                })

                table_id += 1

        # ─────────────────────────────────────────────
        # 2. TEXT (skip table regions)
        # ─────────────────────────────────────────────
        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue

            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if not text:
                        continue

                    bbox = span["bbox"]

                    # Skip text inside tables
                    inside_table = False
                    for tb in table_bboxes:
                        if (
                            bbox[0] >= tb[0] and bbox[2] <= tb[2] and
                            bbox[1] >= tb[1] and bbox[3] <= tb[3]
                        ):
                            inside_table = True
                            break

                    if inside_table:
                        continue

                    tid = f"t_{text_id}"

                    elements.append({
                        "type": "text",
                        "id": tid,
                        "bbox": bbox,
                        "page": page_num,
                        "size": span["size"],
                        "font": span.get("font", ""),
                        "color": span.get("color", 0),
                        "flags": span.get("flags", 0)
                    })

                    text_content[tid] = text
                    text_id += 1

        # ─────────────────────────────────────────────
        # 3. IMAGES (optional but recommended)
        # ─────────────────────────────────────────────
        for img in page.get_images(full=True):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)

            if pix.n > 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)

            img_name = f"img_{img_id}.png"
            img_path = os.path.join(img_dir, img_name)
            pix.save(img_path)

            # NOTE: bbox not directly available → optional improvement later
            elements.append({
                "type": "image",
                "id": f"img_{img_id}",
                "page": page_num,
                "bbox": [0, 0, 100, 100]  # placeholder
            })

            assets[f"img_{img_id}"] = {"path": f"images/{img_name}"}
            img_id += 1

    # ─────────────────────────────────────────────
    # FINAL STRUCTURE
    # ─────────────────────────────────────────────
    structure = {
        "elements": elements,
        "text_content": text_content,
        "assets": assets,
        "pages": pages_meta
    }

    json_path = os.path.join(output_dir, "pdf.json")

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, indent=2, ensure_ascii=False)

    # Translation segments (ONLY text, not tables)
    segments = [
        {"id": tid, "text": text}
        for tid, text in text_content.items()
    ]

    return {
        "json_path": json_path,
        "segments": segments,
        "structure": structure
    }

def extract_docx_segments(input_path: str, output_dir: str) -> dict:
    """Extract DOCX → structured segments with style classification."""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document(input_path)
    elements, text_content, assets = [], {}, {}
    text_id = img_id = 0

    rel_images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_id_str = f"img_{img_id}"
            img_path = os.path.join(output_dir, f"{img_id_str}.png")
            with open(img_path, "wb") as f:
                f.write(rel.target_part.blob)
            rel_images[rel.rId] = img_id_str
            assets[img_id_str] = img_path
            img_id += 1

    for i, para in enumerate(doc.paragraphs):
        pid = f"t_{text_id}"
        style = para.style.name
        # classify element role
        role = "heading" if style.startswith("Heading") else "body"
        elements.append({"type": "paragraph", "id": pid,
                         "index": i, "style": style, "role": role})
        text_content[pid] = para.text
        text_id += 1

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cid = f"t_{text_id}"
                text_content[cid] = cell.text
                row_data.append(cid)
                text_id += 1
            table_data.append(row_data)
        elements.append({"type": "table", "cells": table_data})

    structure = {"elements": elements, "text_content": text_content, "assets": assets}
    json_path = os.path.join(output_dir, "doc.json")
    with open(json_path, "w") as f:
        json.dump(structure, f, indent=2)

    segments = [{"id": tid, "text": text} for tid, text in text_content.items()
                if text.strip()]
    return {"json_path": json_path, "segments": segments, "structure": structure}