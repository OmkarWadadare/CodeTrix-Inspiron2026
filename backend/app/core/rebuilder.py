"""
rebuilder.py

rebuild_pdf_translated(json_path, translations, output_pdf)
rebuild_docx_translated(json_path, translations, output_docx)

`translations` is a dict mapping segment_id → translated text string.
Any segment_id not present in `translations` falls back to its original text.
"""

"""
rebuilder.py
"""

import json
from pathlib import Path
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from PIL import Image

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ─────────────────────────────────────────────
#  PDF REBUILDER
# ─────────────────────────────────────────────

FONT_DIR = Path(__file__).parent / "fonts"

FONTS = {
    "latin": {
        "normal": FONT_DIR / "NotoSans-Regular.ttf",
        "bold": FONT_DIR / "NotoSans-Bold.ttf",
        "italic": FONT_DIR / "NotoSans-Italic.ttf",
        "bold_italic": FONT_DIR / "NotoSans-BoldItalic.ttf",
    },
    "devanagari": {
        "normal": FONT_DIR / "NotoSansDevanagari-Regular.ttf",
        "bold": FONT_DIR / "NotoSansDevanagari-Bold.ttf",
        "italic": FONT_DIR / "NotoSansDevanagari-Italic.ttf",
        "bold_italic": FONT_DIR / "NotoSansDevanagari-BoldItalic.ttf",
    },
    "tamil": {
        "normal": FONT_DIR / "NotoSansTamil-Regular.ttf",
        "bold": FONT_DIR / "NotoSansTamil-Bold.ttf",
        "italic": FONT_DIR / "NotoSansTamil-Italic.ttf",
        "bold_italic": FONT_DIR / "NotoSansTamil-BoldItalic.ttf",
    },
    "arabic": {
        "normal": FONT_DIR / "NotoSansArabic-Regular.ttf",
        "bold": FONT_DIR / "NotoSansArabic-Bold.ttf",
        "italic": FONT_DIR / "NotoSansArabic-Italic.ttf",
        "bold_italic": FONT_DIR / "NotoSansArabic-BoldItalic.ttf",
    },
    "chinese": {
        "normal": FONT_DIR / "NotoSansSC-Regular.ttf",
        "bold": FONT_DIR / "NotoSansSC-Bold.ttf",
        "italic": FONT_DIR / "NotoSansSC-Italic.ttf",
        "bold_italic": FONT_DIR / "NotoSansSC-BoldItalic.ttf",
    },
}

for script, variants in FONTS.items():
    for style, path in variants.items():
        font_name = f"{script}_{style}"
        pdfmetrics.registerFont(TTFont(font_name, str(path)))


def rebuild_pdf_translated(json_path: str, translations: dict, output_pdf: str) -> None:
    json_path = Path(json_path)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    c = canvas.Canvas(str(output_pdf))

    pages_meta   = data["pages"]
    elements     = data["elements"]
    text_content = data["text_content"]
    assets       = data.get("assets", {})

    num_pages = len(pages_meta)

    for i in range(num_pages):
        page_meta = pages_meta[str(i)] if isinstance(pages_meta, dict) else pages_meta[i]
        pw = page_meta["width"]
        ph = page_meta["height"]

        c.setPageSize((pw, ph))

        flow_items = []

        for el in elements:
            if el["page"] != i:
                continue

            if el["type"] == "text":
                raw_text = translations.get(
                    el["id"],
                    text_content.get(el["id"], "")
                ).strip()

                if not raw_text:
                    continue

                x0, y0, x1, y1 = el["bbox"]
                flags = el.get("flags", 0)

                flow_items.append({
                    "type": "text",
                    "text": raw_text,
                    "x0": x0,
                    "y1": y1,
                    "size": el.get("size", 10),
                    "bold": bool(flags & 2),
                    "italic": bool(flags & 1),
                })

            elif el["type"] == "image":
                asset = assets.get(el["id"])
                if not asset:
                    continue

                flow_items.append({
                    "type": "image",
                    "path": json_path.parent / asset["path"],
                    "x0": el["bbox"][0],
                    "y1": el["bbox"][3],
                    "width": el["bbox"][2] - el["bbox"][0],
                    "height": el["bbox"][3] - el["bbox"][1],
                })
            elif el["type"] == "table":
                flow_items.append({
                    "type": "table",
                    "rows": el["rows"],
                    "x0": el["bbox"][0],
                    "y1": el["bbox"][3],
                    "y_anchor": el["bbox"][3],
                    "width": el["bbox"][2] - el["bbox"][0],
                })

        # SORT by reading order
        flow_items.sort(key=lambda b: (-b.get("y_anchor", b["y1"]), b.get("x0", 0)))
        _draw_flow(c, flow_items, pw, ph, text_content, translations)

        c.showPage()

    c.save()
    print(f"[rebuild_pdf_translated] → {output_pdf}")

def _pick_font(text: str, bold: bool = False, italic: bool = False) -> str:
    script = "latin"

    for ch in text:
        code = ord(ch)
        if 0x0900 <= code <= 0x097F:
            script = "devanagari"
            break
        elif 0x0B80 <= code <= 0x0BFF:
            script = "tamil"
            break
        elif 0x0600 <= code <= 0x06FF:
            script = "arabic"
            break
        elif 0x4E00 <= code <= 0x9FFF:
            script = "chinese"
            break

    if bold and italic:
        style = "bold_italic"
    elif bold:
        style = "bold"
    elif italic:
        style = "italic"
    else:
        style = "normal"

    return f"{script}_{style}"
# ─────────────────────────────────────────────
# FLOW ENGINE
# ─────────────────────────────────────────────

def _draw_flow(c, items, page_width, page_height, text_content, translations):
    styles = getSampleStyleSheet()

    LEFT = 40
    RIGHT = 40
    TOP = page_height - 40
    BOTTOM = 40

    usable_width = page_width - LEFT - RIGHT
    y_cursor = TOP

    for item in items:

        # ───────── TEXT ─────────
        if item["type"] == "text":
            text = item["text"]

            style = styles["Normal"]
            style.fontName = _pick_font(text,bold=item.get("bold", False),italic=item.get("italic", False),)
            style.fontSize = item.get("size", 11)
            style.leading = style.fontSize + 2

            # Detect bullets
            if text.strip().startswith(("•", "-", "*")):
                text = f"&bull; {text.lstrip('•-* ').strip()}"
                style.leftIndent = 10

            para = Paragraph(text, style)
            w, h = para.wrap(usable_width, page_height)

            if y_cursor - h < BOTTOM:
                c.showPage()
                y_cursor = TOP

            para.drawOn(c, LEFT, y_cursor - h)
            y_cursor -= (h + 6)

        # ───────── IMAGE ─────────
        elif item["type"] == "image":
            path = item["path"]

            if not path.exists():
                continue

            try:
                w = item["width"]
                h = item["height"]

                # If image too wide, scale proportionally
                if w > usable_width:
                    scale = usable_width / w
                    w *= scale
                    h *= scale

                # Page break if needed
                if y_cursor - h < BOTTOM:
                    c.showPage()
                    y_cursor = TOP

                c.drawImage(
                    str(path),
                    LEFT,
                    y_cursor - h,
                    width=w,
                    height=h,
                    preserveAspectRatio=True,
                    mask='auto'
                )

                y_cursor -= (h + 10)

            except Exception:
                pass
        # ───────── TABLE ─────────
        elif item["type"] == "table":
            data = []
            for row in item["rows"]:
                new_row = []
                for cell in row:
                    if cell is None:
                        new_row.append("")
                    else:
                        if isinstance(cell, dict):
                            seg_id = cell.get("id")
                            flags = cell.get("flags", 0)
                        else:
                            seg_id = cell
                            flags = 0

                        txt = translations.get(
                            seg_id,
                            text_content.get(seg_id, "")
                        )

                        styles = getSampleStyleSheet()
                        cell_style = styles["Normal"]
                        cell_style.fontSize = 9
                        cell_style.leading = 11

                        cell_style.fontName = _pick_font(
                            txt,
                            bold=bool(flags & 2),
                            italic=bool(flags & 1),
                        )

                        para = Paragraph(txt, cell_style)
                        new_row.append(para)
                data.append(new_row)
            # Replace None with empty string
            clean_data = [
                [cell if cell is not None else "" for cell in row]
                for row in data
            ]

            num_cols = max(len(r) for r in data) if data else 1
            col_width = usable_width / num_cols
            col_widths = [col_width] * num_cols
            table = Table(clean_data, colWidths=col_widths)

            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))

            w, h = table.wrap(usable_width, page_height)

            if y_cursor - h < BOTTOM:
                c.showPage()
                y_cursor = TOP

            table.drawOn(c, LEFT, y_cursor - h)
            y_cursor -= (h + 10)

# ─────────────────────────────────────────────
# FONT HELPER
# ─────────────────────────────────────────────

def _select_reportlab_font(bold: bool, italic: bool) -> str:
    if bold and italic:
        return "Helvetica-BoldOblique"
    if bold:
        return "Helvetica-Bold"
    if italic:
        return "Helvetica-Oblique"
    return "Helvetica"

# ─────────────────────────────────────────────
#  DOCX REBUILDER
# ─────────────────────────────────────────────

def rebuild_docx_translated(json_path: str, translations: dict, output_docx: str) -> None:
    from docx import Document
    from docx.shared import Pt
    from pathlib import Path
    import json

    json_path = Path(json_path)

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    elements     = data["elements"]
    text_content = data["text_content"]
    assets       = data.get("assets", {})

    doc = Document()

    def resolve_text(seg_id):
        return translations.get(seg_id, text_content.get(seg_id, "")).strip()

    # ─────────────────────────────
    # FLOW (same concept as PDF)
    # ─────────────────────────────
    for el in elements:

        # ───────── TEXT ─────────
        if el["type"] == "text":
            text = resolve_text(el["id"])
            if not text:
                continue

            p = doc.add_paragraph()
            run = p.add_run(text)

            flags = el.get("flags", 0)
            run.bold   = bool(flags & 2)
            run.italic = bool(flags & 1)

            if el.get("font"):
                run.font.name = el["font"]

            if el.get("size"):
                run.font.size = Pt(el["size"])

        # ───────── IMAGE ─────────
        elif el["type"] == "image":
            asset = assets.get(el["id"])
            if not asset:
                continue

            img_path = json_path.parent / asset["path"]

            if img_path.exists():
                try:
                    doc.add_picture(str(img_path))
                except Exception:
                    pass

        # ───────── TABLE ─────────
        elif el["type"] == "table":
            rows_data = el.get("rows", [])
            if not rows_data:
                continue

            num_rows = len(rows_data)
            num_cols = max(len(r) for r in rows_data)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            set_table_borders(table)

            for r_idx, row in enumerate(rows_data):
                for c_idx, cell in enumerate(row):
                    doc_cell = table.cell(r_idx, c_idx)
                    doc_cell.text = ""

                    p = doc_cell.paragraphs[0]

                    # Each cell can have MULTIPLE segment IDs
                    if isinstance(cell, list):
                        seg_ids = cell
                    else:
                        seg_ids = [cell]

                    for seg_id in seg_ids:
                        if seg_id is None:
                            p.add_run("")
                            continue

                        text = resolve_text(seg_id)
                        if not text:
                            continue

                        run = p.add_run(text)

                        # Optional: style (DOCX doesn't store flags in your new pipeline)
                        # Keeping simple like PDF table rendering

    doc.save(output_docx)
    print(f"[rebuild_docx_translated] → {output_docx}")

def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')   # solid line
        border.set(qn('w:sz'), '8')         # thickness (8 = ~1pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # black

        tblBorders.append(border)

    tblPr.append(tblBorders)