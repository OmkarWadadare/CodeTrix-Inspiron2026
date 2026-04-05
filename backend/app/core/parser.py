import json
import os
import re
import hashlib
import zipfile
import io
from pathlib import Path
from typing import Optional
    
import pypdfium2 as pdfium
import pdfplumber
from docx import Document

def extract_pdf_segments(input_pdf: str, output_dir: str) -> dict:
    """
    Extract all content from a PDF into structured JSON + images folder.

    Parameters
    ----------
    input_pdf : str   Path to the source PDF file.
    output_dir : str  Directory where output is written:
                        <output_dir>/extracted_data.json
                        <output_dir>/images/           (PNG images embedded in the PDF)

    Returns
    -------
    dict  The full extracted data structure (also written to extracted_data.json).
    """
    output_dir = Path(output_dir)
    images_dir = output_dir / "images"
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir.mkdir(parents=True, exist_ok=True)

    doc_data = {
        "source_file": str(Path(input_pdf).resolve()),
        "metadata": {},
        "pages": []
    }

    # ── 1. Metadata via pdfplumber ────────────────────────────────────────
    with pdfplumber.open(input_pdf) as plumber_pdf:
        meta = plumber_pdf.metadata or {}
        doc_data["metadata"] = {k: str(v) for k, v in meta.items()}

        total_pages = len(plumber_pdf.pages)
        doc_data["total_pages"] = total_pages

        for page_num, plumber_page in enumerate(plumber_pdf.pages):
            page_width  = float(plumber_page.width)
            page_height = float(plumber_page.height)

            page_entry = {
                "page_number": page_num + 1,
                "width":  page_width,
                "height": page_height,
                "text_blocks": [],
                "tables": [],
                "images": [],
                "drawings": []
            }

            # ── 2a. Rich character-level text with font/style info ────────
            words = plumber_page.extract_words(
                extra_attrs=["fontname", "size", "stroking_color", "non_stroking_color"]
            )
            page_entry["text_blocks"] = _cluster_words_into_blocks(words, page_height)

            # ── 2b. Tables ────────────────────────────────────────────────
            tables = plumber_page.extract_tables()
            table_bboxes = plumber_page.find_tables()
            for idx, (tbl, tbl_obj) in enumerate(zip(tables, table_bboxes)):
                if not tbl:
                    continue
                bbox = tbl_obj.bbox  # (x0, top, x1, bottom) in pdfplumber coords
                page_entry["tables"].append({
                    "table_index": idx,
                    "bbox": {
                        "x0": bbox[0],
                        "y0": page_height - bbox[3],   # flip to PDF coords (bottom-left origin)
                        "x1": bbox[2],
                        "y1": page_height - bbox[1],
                    },
                    "rows": tbl
                })

            doc_data["pages"].append(page_entry)
            # ── 2c. Hyperlinks (annotations) ───────────────────────────────
            page_entry["links"] = []

            for annot in plumber_page.annots or []:
                uri = annot.get("uri")
                if not uri:
                    continue

                x0 = annot["x0"]
                x1 = annot["x1"]
                top = annot["top"]
                bottom = annot["bottom"]

                page_entry["links"].append({
                    "url": uri,
                    "bbox": {
                        "x0": x0,
                        "y0": page_height - bottom,
                        "x1": x1,
                        "y1": page_height - top
                    }
                })

    # ── 3. Image extraction via pypdfium2 ─────────────────────────────────
    pdfium_doc = pdfium.PdfDocument(input_pdf)

    seen_hashes = {}   # dedup identical embedded images

    for page_num in range(len(pdfium_doc)):
        pdfium_page = pdfium_doc[page_num]
        page_height  = doc_data["pages"][page_num]["height"]

        for obj in pdfium_page.get_objects():
            if not isinstance(obj, pdfium.PdfImage):
                continue

            try:
                bitmap    = obj.get_bitmap()
                pil_image = bitmap.to_pil()

                # Deduplicate by pixel hash
                buf = io.BytesIO()
                pil_image.save(buf, format="PNG")
                img_bytes = buf.getvalue()
                img_hash  = hashlib.md5(img_bytes).hexdigest()[:12]

                if img_hash not in seen_hashes:
                    img_filename = f"page{page_num+1}_img_{img_hash}.png"
                    img_path     = images_dir / img_filename
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    seen_hashes[img_hash] = img_filename
                else:
                    img_filename = seen_hashes[img_hash]

                # Position from matrix
                matrix = obj.get_matrix()
                # PdfMatrix has a,b,c,d,e,f attributes
                x0 = float(matrix.e)
                y0 = float(matrix.f)
                img_w = abs(float(matrix.a)) if matrix.a else pil_image.width
                img_h = abs(float(matrix.d)) if matrix.d else pil_image.height

                doc_data["pages"][page_num]["images"].append({
                    "file": str(Path("images") / img_filename),
                    "hash": img_hash,
                    "pixel_width":  pil_image.width,
                    "pixel_height": pil_image.height,
                    "bbox": {
                        "x0": x0,
                        "y0": y0,
                        "x1": x0 + img_w,
                        "y1": y0 + img_h
                    }
                })
            except Exception as e:
                # Some image objects may not be extractable (masks, etc.)
                pass

    pdfium_doc.close()

    # ── 4. Save JSON ──────────────────────────────────────────────────────
    json_path = output_dir / "extracted_data.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(doc_data, f, indent=2, ensure_ascii=False)

    print(f"[extract_pdf] Done.")
    print(f"  Pages    : {doc_data['total_pages']}")
    print(f"  JSON     : {json_path}")
    print(f"  Images   : {images_dir} ({len(seen_hashes)} unique image(s))")

    # ── 5. Convert to segment-style output (COMPAT LAYER) ───────────────────

    elements = []
    text_content = {}
    assets = {}
    pages_meta = {}

    text_id = 0
    img_id = 0

    for page in doc_data["pages"]:
        page_num = page["page_number"] - 1
        pages_meta[page_num] = {
            "width": page["width"],
            "height": page["height"]
        }

        # ── Text blocks → segments ─────────────────────────────────────────
        for block in page.get("text_blocks", []):
            # ❗ Skip text that belongs to tables
            if _is_inside_table(block, page.get("tables", [])):
                continue
            text = block.get("text", "").strip()
            text = text.replace("'", "").replace('"', "")
            if not text:
                continue

            tid = f"t_{text_id}"
            elements.append({
                "type": "text",
                "id": tid,
                "bbox": [block["x0"], block["y0"], block["x1"], block["y1"]],
                "page": page_num,
                "size": block.get("size", 12),
                "font": block.get("font", ""),
                "color": block.get("color"),
                "flags": (
                    (2 if block.get("bold") else 0) |
                    (1 if block.get("italic") else 0)
                )
            })
            text_content[tid] = text
            text_id += 1

        # ── Images → assets + elements ─────────────────────────────────────
        for img in page.get("images", []):
            iid = f"img_{img_id}"

            elements.append({
                "type": "image",
                "id": iid,
                "bbox": [
                    img["bbox"]["x0"],
                    img["bbox"]["y0"],
                    img["bbox"]["x1"],
                    img["bbox"]["y1"]
                ],
                "page": page_num
            })
            

            assets[iid] = {
                "path": img["file"],
                "width": img.get("pixel_width"),
                "height": img.get("pixel_height")
            }

            img_id += 1

        # ── Tables → elements + segments ─────────────────────────────
        for tbl in page.get("tables", []):
            tid = f"table_{len(elements)}"

            table_rows = []
            
            for row in tbl["rows"]:
                new_row = []
                for cell in row:
                    if not cell or not str(cell).strip():
                        new_row.append(None)
                        continue

                    seg_id = f"t_{text_id}"
                    cell_text = str(cell).replace("'", "").replace('"', "")
                    text_content[seg_id] = cell_text
                    new_row.append(seg_id)
                    text_id += 1

                table_rows.append(new_row)

            elements.append({
                "type": "table",
                "id": tid,
                "page": page_num,
                "bbox": [
                    tbl["bbox"]["x0"],
                    tbl["bbox"]["y0"],
                    tbl["bbox"]["x1"],
                    tbl["bbox"]["y1"]
                ],
                "rows": table_rows   # ← NOW stores segment IDs
            })
    # ── Build final structure ──────────────────────────────────────────────
    structure = {
        "elements": elements,
        "text_content": text_content,
        "assets": assets,
        "pages": pages_meta
    }

    # Save in SAME format as extract_pdf_segments
    json_path = output_dir / "pdf.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, indent=2, ensure_ascii=False)

    # Flat segments (for translation)
    segments = [{"id": tid, "text": text} for tid, text in text_content.items()]

    print(f"[extract_pdf] Segment JSON → {json_path}")

    return {
        "json_path": str(json_path),
        "segments": segments,
        "structure": structure
    }


def _cluster_words_into_blocks(words: list, page_height: float) -> list:
    """
    Group individual words (with font metadata) into logical text blocks
    by proximity and shared style properties.
    """
    if not words:
        return []

    # Sort top-to-bottom, left-to-right
    words = sorted(words, key=lambda w: (round(w["top"] / 4) * 4, w["x0"]))

    blocks = []
    current_block = None

    for w in words:
        text  = w.get("text", "")
        font  = w.get("fontname", "")
        size  = float(w.get("size", 12) or 12)
        x0    = float(w["x0"])
        top   = float(w["top"])
        x1    = float(w["x1"])
        bot   = float(w["bottom"])

        # Detect style from font name heuristics
        fn_lower = font.lower()
        is_bold   = any(kw in fn_lower for kw in ["bold", "black", "heavy", "demi", "semibold"])
        is_italic = any(kw in fn_lower for kw in ["italic", "oblique", "slant"])

        # Determine semantic role from size
        role = _classify_text_role(text, size, is_bold)

        # Colour
        color = w.get("non_stroking_color")
        color_hex = _color_to_hex(color)

        word_info = {
            "text":      text,
            "x0":        x0,
            "y0":        page_height - bot,   # flip to PDF coords
            "x1":        x1,
            "y1":        page_height - top,
            "font":      font,
            "size":      size,
            "bold":      is_bold,
            "italic":    is_italic,
            "color":     color_hex,
            "role":      role
        }

        # Try to append to current block if same line + style
        if current_block and _same_block(current_block, word_info):
            current_block["words"].append(word_info)
            current_block["x1"]    = max(current_block["x1"], x1)
            current_block["y0"]    = min(current_block["y0"], page_height - bot)
            current_block["y1"]    = max(current_block["y1"], page_height - top)
            current_block["text"] += " " + text
        else:
            if current_block:
                blocks.append(current_block)
            current_block = {
                "text":   text,
                "x0":     x0,
                "y0":     page_height - bot,
                "x1":     x1,
                "y1":     page_height - top,
                "font":   font,
                "size":   size,
                "bold":   is_bold,
                "italic": is_italic,
                "color":  color_hex,
                "role":   role,
                "words":  [word_info]
            }

    if current_block:
        blocks.append(current_block)

    return blocks


def _same_block(block: dict, word: dict) -> bool:
    """Heuristic: same block if same line (≤4pt vertical diff) and same font/size."""
    v_diff = abs(block["y1"] - word["y1"])
    same_line  = v_diff < max(4, block["size"] * 0.5)
    same_font  = block["font"] == word["font"]
    same_size  = abs(block["size"] - word["size"]) < 0.5
    close_x    = word["x0"] - block["x1"] < block["size"] * 3  # gap < 3 em-widths
    return same_line and same_font and same_size and close_x


def _classify_text_role(text: str, size: float, bold: bool) -> str:
    """Classify text as heading level, list item, or body."""
    stripped = text.strip()
    if size >= 20:
        return "h1"
    if size >= 16:
        return "h2"
    if size >= 13 and bold:
        return "h3"
    # List item heuristics
    if re.match(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s", stripped):
        return "list_item"
    if re.match(r"^\d+[\.\)]\s", stripped):
        return "ordered_list_item"
    return "body"


def _color_to_hex(color) -> Optional[str]:
    """Convert a pdfplumber color value to #RRGGBB string."""
    if color is None:
        return None
    if isinstance(color, (list, tuple)):
        if len(color) == 3:
            r, g, b = [int(c * 255) if c <= 1.0 else int(c) for c in color]
            return f"#{r:02x}{g:02x}{b:02x}"
        if len(color) == 4:  # CMYK
            c, m, y, k = color
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return f"#{r:02x}{g:02x}{b:02x}"
    if isinstance(color, (int, float)):
        v = int(color * 255) if color <= 1.0 else int(color)
        return f"#{v:02x}{v:02x}{v:02x}"
    return None

def _is_inside_table(block, tables):
    bx0, by0, bx1, by1 = block["x0"], block["y0"], block["x1"], block["y1"]

    for tbl in tables:
        tx0 = tbl["bbox"]["x0"]
        ty0 = tbl["bbox"]["y0"]
        tx1 = tbl["bbox"]["x1"]
        ty1 = tbl["bbox"]["y1"]

        # Check overlap
        if not (bx1 < tx0 or bx0 > tx1 or by1 < ty0 or by0 > ty1):
            return True
    return False

# ─────────────────────────────────────────────
#  DOCX PARSER
# ─────────────────────────────────────────────

def extract_docx_segments(input_path: str, output_dir: str) -> dict:
    
    output_dir = Path(output_dir)
    img_dir = output_dir / "images"
    output_dir.mkdir(parents=True, exist_ok=True)
    img_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(input_path)

    elements = []
    text_content = {}
    assets = {}

    text_id = 0
    img_id = 0

    def clean_text(text: str) -> str:
        return text.replace("'", "").replace('"', "").strip()

    # ── Extract images ─────────────────────────────
    rel_to_img = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            iid = f"img_{img_id}"
            img_path = img_dir / f"{iid}.png"

            with open(img_path, "wb") as f:
                f.write(rel.target_part.blob)

            rel_to_img[rel.rId] = iid
            assets[iid] = {
                "path": str(Path("images") / f"{iid}.png")
            }
            img_id += 1

    # ── Paragraphs ────────────────────────────────
    for para in doc.paragraphs:
        for run in para.runs:

            # Handle images inside runs
            blips = run._element.xpath('.//a:blip')
            if blips:
                rId = blips[0].attrib.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if rId in rel_to_img:
                    iid = rel_to_img[rId]

                    elements.append({
                        "type": "image",
                        "id": iid,
                        "page": 0,   # DOCX has no pages
                        "bbox": None
                    })

            text = clean_text(run.text)
            if not text:
                continue

            tid = f"t_{text_id}"

            elements.append({
                "type": "text",
                "id": tid,
                "page": 0,
                "bbox": None,
                "size": run.font.size.pt if run.font.size else None,
                "font": run.font.name,
                "color": None,
                "flags": (
                    (2 if run.bold else 0) |
                    (1 if run.italic else 0)
                )
            })

            text_content[tid] = text
            text_id += 1

    # ── Tables ────────────────────────────────────
    for table in doc.tables:
        table_rows = []

        for row in table.rows:
            new_row = []

            for cell in row.cells:
                cell_segments = []

                for para in cell.paragraphs:
                    for run in para.runs:
                        text = clean_text(run.text)

                        if not text:
                            cell_segments.append(None)
                            continue

                        seg_id = f"t_{text_id}"
                        text_content[seg_id] = text
                        cell_segments.append(seg_id)
                        text_id += 1

                new_row.append(cell_segments if cell_segments else [None])

            table_rows.append(new_row)

        elements.append({
            "type": "table",
            "id": f"table_{len(elements)}",
            "page": 0,
            "bbox": None,
            "rows": table_rows
        })

    # ── Final structure (MATCHES PDF OUTPUT) ──────
    structure = {
        "elements": elements,
        "text_content": text_content,
        "assets": assets,
        "pages": {
            0: {"width": None, "height": None}
        }
    }

    json_path = output_dir / "doc.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, indent=2, ensure_ascii=False)

    segments = [{"id": k, "text": v} for k, v in text_content.items()]

    return {
        "json_path": str(json_path),
        "segments": segments,
        "structure": structure
    }

# ─────────────────────────────────────────────
#  INTERNAL HELPERS (shared)
# ─────────────────────────────────────────────

def _cluster_words_into_blocks(words: list, page_height: float) -> list:
    if not words:
        return []

    words = sorted(words, key=lambda w: (round(w["top"] / 4) * 4, w["x0"]))
    blocks = []
    current_block = None

    for w in words:
        text = w.get("text", "")
        font = w.get("fontname", "")
        size = float(w.get("size", 12) or 12)
        x0 = float(w["x0"])
        top = float(w["top"])
        x1 = float(w["x1"])
        bot = float(w["bottom"])

        fn_lower = font.lower()
        is_bold   = any(kw in fn_lower for kw in ["bold", "black", "heavy", "demi", "semibold"])
        is_italic = any(kw in fn_lower for kw in ["italic", "oblique", "slant"])
        role = _classify_text_role(text, size, is_bold)
        color_hex = _color_to_hex(w.get("non_stroking_color"))

        word_info = {
            "text": text, "x0": x0, "y0": page_height - bot,
            "x1": x1, "y1": page_height - top,
            "font": font, "size": size, "bold": is_bold,
            "italic": is_italic, "color": color_hex, "role": role,
        }

        if current_block and _same_block(current_block, word_info):
            current_block["words"].append(word_info)
            current_block["x1"]    = max(current_block["x1"], x1)
            current_block["y0"]    = min(current_block["y0"], page_height - bot)
            current_block["y1"]    = max(current_block["y1"], page_height - top)
            current_block["text"] += " " + text
        else:
            if current_block:
                blocks.append(current_block)
            current_block = {
                "text": text, "x0": x0, "y0": page_height - bot,
                "x1": x1, "y1": page_height - top,
                "font": font, "size": size, "bold": is_bold,
                "italic": is_italic, "color": color_hex, "role": role,
                "words": [word_info],
            }

    if current_block:
        blocks.append(current_block)
    return blocks


def _same_block(block: dict, word: dict) -> bool:
    v_diff = abs(block["y1"] - word["y1"])
    same_line = v_diff < max(4, block["size"] * 0.5)
    same_font = block["font"] == word["font"]
    same_size = abs(block["size"] - word["size"]) < 0.5
    close_x   = word["x0"] - block["x1"] < block["size"] * 3
    return same_line and same_font and same_size and close_x


def _classify_text_role(text: str, size: float, bold: bool) -> str:
    stripped = text.strip()
    if size >= 20:
        return "h1"
    if size >= 16:
        return "h2"
    if size >= 13 and bold:
        return "h3"
    if re.match(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s", stripped):
        return "list_item"
    if re.match(r"^\d+[\.\)]\s", stripped):
        return "ordered_list_item"
    return "body"


def _color_to_hex(color) -> Optional[str]:
    if color is None:
        return None
    if isinstance(color, (list, tuple)):
        if len(color) == 3:
            r, g, b = [int(c * 255) if c <= 1.0 else int(c) for c in color]
            return f"#{r:02x}{g:02x}{b:02x}"
        if len(color) == 4:
            c, m, y, k = color
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return f"#{r:02x}{g:02x}{b:02x}"
    if isinstance(color, (int, float)):
        v = int(color * 255) if color <= 1.0 else int(color)
        return f"#{v:02x}{v:02x}{v:02x}"
    return None