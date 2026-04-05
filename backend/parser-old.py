import json
import os
import re
import hashlib
import zipfile
import io
from pathlib import Path
from typing import Optional

import pypdfium2 as pdfium
import pdfplumber
from PIL import Image
from docx import Document

def extract_pdf_segments(input_pdf: str, output_dir: str) -> dict:
    """
    Extract all content from a PDF into structured JSON + images folder.

    Parameters
    ----------
    input_pdf : str   Path to the source PDF file.
    output_dir : str  Directory where output is written:
                        <output_dir>/extracted_data.json
                        <output_dir>/images/           (PNG images embedded in the PDF)

    Returns
    -------
    dict  The full extracted data structure (also written to extracted_data.json).
    """
    output_dir = Path(output_dir)
    images_dir = output_dir / "images"
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir.mkdir(parents=True, exist_ok=True)

    doc_data = {
        "source_file": str(Path(input_pdf).resolve()),
        "metadata": {},
        "pages": []
    }

    # ── 1. Metadata via pdfplumber ────────────────────────────────────────
    with pdfplumber.open(input_pdf) as plumber_pdf:
        meta = plumber_pdf.metadata or {}
        doc_data["metadata"] = {k: str(v) for k, v in meta.items()}

        total_pages = len(plumber_pdf.pages)
        doc_data["total_pages"] = total_pages

        for page_num, plumber_page in enumerate(plumber_pdf.pages):
            page_width  = float(plumber_page.width)
            page_height = float(plumber_page.height)

            page_entry = {
                "page_number": page_num + 1,
                "width":  page_width,
                "height": page_height,
                "text_blocks": [],
                "tables": [],
                "images": [],
                "drawings": []
            }

            # ── 2a. Rich character-level text with font/style info ────────
            words = plumber_page.extract_words(
                extra_attrs=["fontname", "size", "stroking_color", "non_stroking_color"]
            )
            page_entry["text_blocks"] = _cluster_words_into_blocks(words, page_height)

            # ── 2b. Tables ────────────────────────────────────────────────
            tables = plumber_page.extract_tables()
            table_bboxes = plumber_page.find_tables()
            for idx, (tbl, tbl_obj) in enumerate(zip(tables, table_bboxes)):
                if not tbl:
                    continue
                bbox = tbl_obj.bbox  # (x0, top, x1, bottom) in pdfplumber coords
                page_entry["tables"].append({
                    "table_index": idx,
                    "bbox": {
                        "x0": bbox[0],
                        "y0": page_height - bbox[3],   # flip to PDF coords (bottom-left origin)
                        "x1": bbox[2],
                        "y1": page_height - bbox[1],
                    },
                    "rows": tbl
                })

            doc_data["pages"].append(page_entry)
            # ── 2c. Hyperlinks (annotations) ───────────────────────────────
            page_entry["links"] = []

            for annot in plumber_page.annots or []:
                uri = annot.get("uri")
                if not uri:
                    continue

                x0 = annot["x0"]
                x1 = annot["x1"]
                top = annot["top"]
                bottom = annot["bottom"]

                page_entry["links"].append({
                    "url": uri,
                    "bbox": {
                        "x0": x0,
                        "y0": page_height - bottom,
                        "x1": x1,
                        "y1": page_height - top
                    }
                })

    # ── 3. Image extraction via pypdfium2 ─────────────────────────────────
    pdfium_doc = pdfium.PdfDocument(input_pdf)

    seen_hashes = {}   # dedup identical embedded images

    for page_num in range(len(pdfium_doc)):
        pdfium_page = pdfium_doc[page_num]
        page_height  = doc_data["pages"][page_num]["height"]

        for obj in pdfium_page.get_objects():
            if not isinstance(obj, pdfium.PdfImage):
                continue

            try:
                bitmap    = obj.get_bitmap()
                pil_image = bitmap.to_pil()

                # Deduplicate by pixel hash
                buf = io.BytesIO()
                pil_image.save(buf, format="PNG")
                img_bytes = buf.getvalue()
                img_hash  = hashlib.md5(img_bytes).hexdigest()[:12]

                if img_hash not in seen_hashes:
                    img_filename = f"page{page_num+1}_img_{img_hash}.png"
                    img_path     = images_dir / img_filename
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    seen_hashes[img_hash] = img_filename
                else:
                    img_filename = seen_hashes[img_hash]

                # Position from matrix
                matrix = obj.get_matrix()
                # PdfMatrix has a,b,c,d,e,f attributes
                x0 = float(matrix.e)
                y0 = float(matrix.f)
                img_w = abs(float(matrix.a)) if matrix.a else pil_image.width
                img_h = abs(float(matrix.d)) if matrix.d else pil_image.height

                doc_data["pages"][page_num]["images"].append({
                    "file": str(Path("images") / img_filename),
                    "hash": img_hash,
                    "pixel_width":  pil_image.width,
                    "pixel_height": pil_image.height,
                    "bbox": {
                        "x0": x0,
                        "y0": y0,
                        "x1": x0 + img_w,
                        "y1": y0 + img_h
                    }
                })
            except Exception as e:
                # Some image objects may not be extractable (masks, etc.)
                pass

    pdfium_doc.close()

    # ── 4. Save JSON ──────────────────────────────────────────────────────
    json_path = output_dir / "extracted_data.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(doc_data, f, indent=2, ensure_ascii=False)

    print(f"[extract_pdf] Done.")
    print(f"  Pages    : {doc_data['total_pages']}")
    print(f"  JSON     : {json_path}")
    print(f"  Images   : {images_dir} ({len(seen_hashes)} unique image(s))")

    # ── 5. Convert to segment-style output (COMPAT LAYER) ───────────────────

    elements = []
    text_content = {}
    assets = {}
    pages_meta = {}

    text_id = 0
    img_id = 0

    for page in doc_data["pages"]:
        page_num = page["page_number"] - 1
        pages_meta[page_num] = {
            "width": page["width"],
            "height": page["height"]
        }

        # ── Text blocks → segments ─────────────────────────────────────────
        for block in page.get("text_blocks", []):
            text = block.get("text", "").strip()
            if not text:
                continue

            tid = f"t_{text_id}"
            elements.append({
                "type": "text",
                "id": tid,
                "bbox": [block["x0"], block["y0"], block["x1"], block["y1"]],
                "page": page_num,
                "size": block.get("size", 12),
                "font": block.get("font", ""),
                "color": block.get("color"),
                "flags": (
                    (2 if block.get("bold") else 0) |
                    (1 if block.get("italic") else 0)
                )
            })
            text_content[tid] = text
            text_id += 1

        # ── Images → assets + elements ─────────────────────────────────────
        for img in page.get("images", []):
            iid = f"img_{img_id}"

            elements.append({
                "type": "image",
                "id": iid,
                "bbox": [
                    img["bbox"]["x0"],
                    img["bbox"]["y0"],
                    img["bbox"]["x1"],
                    img["bbox"]["y1"]
                ],
                "page": page_num
            })

            assets[iid] = {
                "path": img["file"],
                "width": img.get("pixel_width"),
                "height": img.get("pixel_height")
            }

            img_id += 1

    # ── Build final structure ──────────────────────────────────────────────
    structure = {
        "elements": elements,
        "text_content": text_content,
        "assets": assets,
        "pages": pages_meta
    }

    # Save in SAME format as extract_pdf_segments
    json_path = output_dir / "pdf.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(structure, f, indent=2, ensure_ascii=False)

    # Flat segments (for translation)
    segments = [{"id": tid, "text": text} for tid, text in text_content.items()]

    print(f"[extract_pdf] Segment JSON → {json_path}")

    return {
        "json_path": str(json_path),
        "segments": segments,
        "structure": structure
    }


def _cluster_words_into_blocks(words: list, page_height: float) -> list:
    """
    Group individual words (with font metadata) into logical text blocks
    by proximity and shared style properties.
    """
    if not words:
        return []

    # Sort top-to-bottom, left-to-right
    words = sorted(words, key=lambda w: (round(w["top"] / 4) * 4, w["x0"]))

    blocks = []
    current_block = None

    for w in words:
        text  = w.get("text", "")
        font  = w.get("fontname", "")
        size  = float(w.get("size", 12) or 12)
        x0    = float(w["x0"])
        top   = float(w["top"])
        x1    = float(w["x1"])
        bot   = float(w["bottom"])

        # Detect style from font name heuristics
        fn_lower = font.lower()
        is_bold   = any(kw in fn_lower for kw in ["bold", "black", "heavy", "demi", "semibold"])
        is_italic = any(kw in fn_lower for kw in ["italic", "oblique", "slant"])

        # Determine semantic role from size
        role = _classify_text_role(text, size, is_bold)

        # Colour
        color = w.get("non_stroking_color")
        color_hex = _color_to_hex(color)

        word_info = {
            "text":      text,
            "x0":        x0,
            "y0":        page_height - bot,   # flip to PDF coords
            "x1":        x1,
            "y1":        page_height - top,
            "font":      font,
            "size":      size,
            "bold":      is_bold,
            "italic":    is_italic,
            "color":     color_hex,
            "role":      role
        }

        # Try to append to current block if same line + style
        if current_block and _same_block(current_block, word_info):
            current_block["words"].append(word_info)
            current_block["x1"]    = max(current_block["x1"], x1)
            current_block["y0"]    = min(current_block["y0"], page_height - bot)
            current_block["y1"]    = max(current_block["y1"], page_height - top)
            current_block["text"] += " " + text
        else:
            if current_block:
                blocks.append(current_block)
            current_block = {
                "text":   text,
                "x0":     x0,
                "y0":     page_height - bot,
                "x1":     x1,
                "y1":     page_height - top,
                "font":   font,
                "size":   size,
                "bold":   is_bold,
                "italic": is_italic,
                "color":  color_hex,
                "role":   role,
                "words":  [word_info]
            }

    if current_block:
        blocks.append(current_block)

    return blocks


def _same_block(block: dict, word: dict) -> bool:
    """Heuristic: same block if same line (≤4pt vertical diff) and same font/size."""
    v_diff = abs(block["y1"] - word["y1"])
    same_line  = v_diff < max(4, block["size"] * 0.5)
    same_font  = block["font"] == word["font"]
    same_size  = abs(block["size"] - word["size"]) < 0.5
    close_x    = word["x0"] - block["x1"] < block["size"] * 3  # gap < 3 em-widths
    return same_line and same_font and same_size and close_x


def _classify_text_role(text: str, size: float, bold: bool) -> str:
    """Classify text as heading level, list item, or body."""
    stripped = text.strip()
    if size >= 20:
        return "h1"
    if size >= 16:
        return "h2"
    if size >= 13 and bold:
        return "h3"
    # List item heuristics
    if re.match(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s", stripped):
        return "list_item"
    if re.match(r"^\d+[\.\)]\s", stripped):
        return "ordered_list_item"
    return "body"


def _color_to_hex(color) -> Optional[str]:
    """Convert a pdfplumber color value to #RRGGBB string."""
    if color is None:
        return None
    if isinstance(color, (list, tuple)):
        if len(color) == 3:
            r, g, b = [int(c * 255) if c <= 1.0 else int(c) for c in color]
            return f"#{r:02x}{g:02x}{b:02x}"
        if len(color) == 4:  # CMYK
            c, m, y, k = color
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return f"#{r:02x}{g:02x}{b:02x}"
    if isinstance(color, (int, float)):
        v = int(color * 255) if color <= 1.0 else int(color)
        return f"#{v:02x}{v:02x}{v:02x}"
    return None

# ─────────────────────────────────────────────
#  DOCX PARSER
# ─────────────────────────────────────────────

def extract_docx_segments(input_path: str, output_dir: str) -> dict:
    """
    Extract translatable text segments and images from a DOCX.

    Returns
    -------
    dict with keys:
        source_file  : str
        segments     : list[dict]   flat list of translatable text segments
        elements     : list[dict]   ordered document structure (for rebuilding)
        images_dir   : str

    Each segment dict:
        segment_id   : str
        element_ref  : str          links back to the element that owns it
        run_index    : int | None   which run within a paragraph (-1 = table cell)
        text         : str
        bold         : bool
        italic       : bool
        underline    : bool
        font_name    : str | None
        font_size    : float | None
        style        : str | None   paragraph style name
        in_table     : bool
        table_index  : int | None
        row          : int | None
        col          : int | None
    """
    output_dir = Path(output_dir)
    img_dir = output_dir / "images"
    font_dir = output_dir / "fonts"
    output_dir.mkdir(parents=True, exist_ok=True)
    img_dir.mkdir(parents=True, exist_ok=True)
    font_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(input_path)

    doc_data = {
        "source_file": str(Path(input_path).resolve()),
        "segments": [],
        "elements": [],   # ordered rebuild blueprint
        "images_dir": str(img_dir.resolve()),
        "assets": {},     # image ref → absolute path
        "fonts": {},
    }

    seg_counter = [0]

    def new_seg_id():
        seg_counter[0] += 1
        return f"seg_{seg_counter[0]}"

    # ── Fonts ─────────────────────────────────────────────────────────────
    with zipfile.ZipFile(input_path, "r") as z:
        for file in z.namelist():
            if file.startswith("word/fonts/"):
                font_name = os.path.basename(file)
                font_path = str(font_dir / font_name)
                with open(font_path, "wb") as f:
                    f.write(z.read(file))
                doc_data["fonts"][font_name] = font_path

    # ── Images ────────────────────────────────────────────────────────────
    rel_to_ref = {}
    img_counter = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            ref = f"img_{img_counter}"
            img_path = str(img_dir / f"{ref}.png")
            with open(img_path, "wb") as f:
                f.write(rel.target_part.blob)
            rel_to_ref[rel.rId] = ref
            doc_data["assets"][ref] = img_path
            img_counter += 1

    # ── Paragraphs ────────────────────────────────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):
        elem_id = f"para_{para_idx}"
        elem = {
            "type": "paragraph",
            "id": elem_id,
            "style": para.style.name,
            "segment_ids": [],
            "image_refs": [],
        }

        for run_idx, run in enumerate(para.runs):
            # Check if run contains an image
            blips = run._element.xpath('.//a:blip')
            if blips:
                rId = blips[0].attrib.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if rId in rel_to_ref:
                    elem["image_refs"].append(rel_to_ref[rId])

            text = run.text
            if not text or not text.strip():
                # Keep empty runs in element for fidelity but no segment
                elem["segment_ids"].append(None)
                continue

            seg_id = new_seg_id()
            seg = {
                "segment_id": seg_id,
                "element_ref": elem_id,
                "run_index": run_idx,
                "text": text,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "underline": bool(run.underline),
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "style": para.style.name,
                "in_table": False,
                "table_index": None,
                "row": None,
                "col": None,
            }
            doc_data["segments"].append(seg)
            elem["segment_ids"].append(seg_id)

        doc_data["elements"].append(elem)

    # ── Tables ────────────────────────────────────────────────────────────
    for tbl_idx, table in enumerate(doc.tables):
        elem_id = f"table_{tbl_idx}"
        elem = {
            "type": "table",
            "id": elem_id,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "cells": [],   # list of rows; each row = list of {segment_ids, style}
        }

        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                cell_seg_ids = []
                cell_style = None

                for para in cell.paragraphs:
                    if cell_style is None:
                        cell_style = para.style.name

                    for run_idx, run in enumerate(para.runs):
                        text = run.text
                        if not text or not text.strip():
                            cell_seg_ids.append(None)
                            continue
                        seg_id = new_seg_id()
                        seg = {
                            "segment_id": seg_id,
                            "element_ref": elem_id,
                            "run_index": run_idx,
                            "text": text,
                            "bold": bool(run.bold),
                            "italic": bool(run.italic),
                            "underline": bool(run.underline),
                            "font_name": run.font.name,
                            "font_size": run.font.size.pt if run.font.size else None,
                            "style": cell_style,
                            "in_table": True,
                            "table_index": tbl_idx,
                            "row": r_idx,
                            "col": c_idx,
                        }
                        doc_data["segments"].append(seg)
                        cell_seg_ids.append(seg_id)

                row_data.append({"segment_ids": cell_seg_ids, "style": cell_style})

            elem["cells"].append(row_data)

        doc_data["elements"].append(elem)

    json_path = output_dir / "segments.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(doc_data, f, indent=2, ensure_ascii=False)

    print(f"[extract_docx_segments] {len(doc_data['segments'])} segments → {json_path}")
    return doc_data


# ─────────────────────────────────────────────
#  INTERNAL HELPERS (shared)
# ─────────────────────────────────────────────

def _cluster_words_into_blocks(words: list, page_height: float) -> list:
    if not words:
        return []

    words = sorted(words, key=lambda w: (round(w["top"] / 4) * 4, w["x0"]))
    blocks = []
    current_block = None

    for w in words:
        text = w.get("text", "")
        font = w.get("fontname", "")
        size = float(w.get("size", 12) or 12)
        x0 = float(w["x0"])
        top = float(w["top"])
        x1 = float(w["x1"])
        bot = float(w["bottom"])

        fn_lower = font.lower()
        is_bold   = any(kw in fn_lower for kw in ["bold", "black", "heavy", "demi", "semibold"])
        is_italic = any(kw in fn_lower for kw in ["italic", "oblique", "slant"])
        role = _classify_text_role(text, size, is_bold)
        color_hex = _color_to_hex(w.get("non_stroking_color"))

        word_info = {
            "text": text, "x0": x0, "y0": page_height - bot,
            "x1": x1, "y1": page_height - top,
            "font": font, "size": size, "bold": is_bold,
            "italic": is_italic, "color": color_hex, "role": role,
        }

        if current_block and _same_block(current_block, word_info):
            current_block["words"].append(word_info)
            current_block["x1"]    = max(current_block["x1"], x1)
            current_block["y0"]    = min(current_block["y0"], page_height - bot)
            current_block["y1"]    = max(current_block["y1"], page_height - top)
            current_block["text"] += " " + text
        else:
            if current_block:
                blocks.append(current_block)
            current_block = {
                "text": text, "x0": x0, "y0": page_height - bot,
                "x1": x1, "y1": page_height - top,
                "font": font, "size": size, "bold": is_bold,
                "italic": is_italic, "color": color_hex, "role": role,
                "words": [word_info],
            }

    if current_block:
        blocks.append(current_block)
    return blocks


def _same_block(block: dict, word: dict) -> bool:
    v_diff = abs(block["y1"] - word["y1"])
    same_line = v_diff < max(4, block["size"] * 0.5)
    same_font = block["font"] == word["font"]
    same_size = abs(block["size"] - word["size"]) < 0.5
    close_x   = word["x0"] - block["x1"] < block["size"] * 3
    return same_line and same_font and same_size and close_x


def _classify_text_role(text: str, size: float, bold: bool) -> str:
    stripped = text.strip()
    if size >= 20:
        return "h1"
    if size >= 16:
        return "h2"
    if size >= 13 and bold:
        return "h3"
    if re.match(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s", stripped):
        return "list_item"
    if re.match(r"^\d+[\.\)]\s", stripped):
        return "ordered_list_item"
    return "body"


def _color_to_hex(color) -> Optional[str]:
    if color is None:
        return None
    if isinstance(color, (list, tuple)):
        if len(color) == 3:
            r, g, b = [int(c * 255) if c <= 1.0 else int(c) for c in color]
            return f"#{r:02x}{g:02x}{b:02x}"
        if len(color) == 4:
            c, m, y, k = color
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return f"#{r:02x}{g:02x}{b:02x}"
    if isinstance(color, (int, float)):
        v = int(color * 255) if color <= 1.0 else int(color)
        return f"#{v:02x}{v:02x}{v:02x}"
    return None